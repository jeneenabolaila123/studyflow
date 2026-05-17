import hashlib
import logging
import os
import re
import tempfile

import docx
import fitz
from app.crud import get_document_by_hash
from sqlalchemy.orm import Session
from app.vectorize import (
    extract_chapter_contents_from_bytes,
    extract_chapter_headings_with_page_numbers_from_bytes,
)

logger = logging.getLogger(__name__)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def hash_content(content):
    if isinstance(content, str):
        content = content.encode("utf-8")
    return hashlib.sha256(content).hexdigest()


def extract_author_from_docx(binary_data: bytes) -> str:
    """Extract author from DOCX core properties, falling back to 'Unknown Author'."""
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(binary_data)
            tmp_path = tmp.name

        doc = docx.Document(tmp_path)
        author = doc.core_properties.author
        return author if author else "Unknown Author"

    except Exception:
        return "Unknown Author"

    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


def extract_author_from_pdf(binary_data: bytes) -> str:
    """Extract author from PDF metadata, falling back to 'Unknown Author'."""
    try:
        pdf = fitz.open(stream=binary_data, filetype="pdf")
        metadata = pdf.metadata
        return metadata.get("author") or "Unknown Author"
    except Exception:
        return "Unknown Author"


def extract_title_from_filename(filename: str) -> str:
    """Extract a human-readable title from the uploaded filename."""
    name = os.path.splitext(os.path.basename(filename))[0]
    return name.replace("_", " ").replace("-", " ").strip() or "Unknown Title"


def save_file(doc_hash, file_ext, content):
    filename = f"{doc_hash}.{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    with open(file_path, "wb") as f:
        f.write(content)


async def process_file(file, db: Session):
    content = await file.read()
    file_ext = file.filename.split(".")[-1].lower()

    if file_ext == "pdf":
        doc_author = extract_author_from_pdf(content)
    elif file_ext == "docx":
        doc_author = extract_author_from_docx(content)
    else:
        raise ValueError("Unsupported file type")

    doc_title = extract_title_from_filename(file.filename)
    doc_hash = hash_content(content)

    # Use content hash for reliable duplicate detection.
    existing_doc = get_document_by_hash(db, doc_hash)

    if existing_doc:
        existing_chapters = list(getattr(existing_doc, "chapters", []) or [])

        # Normal case: document already exists AND has chapters.
        if existing_chapters:
            logger.info("Document already exists with chapters: id=%s", existing_doc.id)
            return {
                "document_id": existing_doc.id,
                "already_exists": True,
                "document_hash": existing_doc.document_hash,
                "title": existing_doc.title,
                "author": existing_doc.author,
                "file_type": existing_doc.file_type,
                "chapters": existing_chapters,
            }

        # Important fix:
        # Sometimes a document was saved before but chapter extraction returned [].
        # If we return it as "existing", Swagger shows chapters: [] forever.
        # Delete only this broken no-chapter record and reprocess the same upload.
        logger.warning(
            "Document exists but has no chapters. Reprocessing stale document id=%s",
            existing_doc.id,
        )
        try:
            db.delete(existing_doc)
            db.commit()
        except Exception:
            db.rollback()
            logger.exception("Failed to delete stale no-chapter document")
            raise

    chapters = _extract_chapters_with_fallback(
        binary_data=content,
        file_ext=file_ext,
        filename=file.filename,
    )

    logger.info("Extracted %d chapters from %s", len(chapters), file.filename)

    for ch in chapters:
        ch["hash"] = hash_content(ch["content"])

    save_file(doc_hash, file_ext, content)

    return {
        "already_exists": False,
        "document_hash": doc_hash,
        "title": doc_title,
        "author": doc_author,
        "file_type": file_ext,
        "chapters": chapters,
    }


def _extract_chapters_with_fallback(binary_data: bytes, file_ext: str, filename: str):
    """
    Keep the old vectorize-based chapter extraction first.
    If it returns no chapters, fall back safely:
    1. Extract plain text.
    2. Try splitting by Chapter/Lecture/Lesson/Unit/Module/Week/Section.
    3. If still nothing, create ONE chapter from the whole PDF/DOCX text.

    This fixes lecture-style PDFs like "Lecture 10" without breaking Chapter PDFs.
    """
    chapters = []

    try:
        chapter_headings = extract_chapter_headings_with_page_numbers_from_bytes(binary_data)
        chapters = extract_chapter_contents_from_bytes(binary_data, chapter_headings) or []
    except Exception:
        logger.exception("Vector chapter extraction failed; using text fallback.")
        chapters = []

    chapters = _clean_chapters(chapters)

    if chapters:
        return chapters

    if file_ext == "pdf":
        text = extract_text_from_pdf(binary_data)
    elif file_ext == "docx":
        text = extract_text_from_docx(binary_data)
    else:
        text = ""

    text = _clean_extracted_text(text)

    if not text:
        return []

    chapters = split_into_chapters(text)
    chapters = _clean_chapters(chapters)

    if chapters:
        return chapters

    # Final fallback for slide decks / lecture notes that do not contain "Chapter".
    title = _guess_single_chapter_title(text, filename)
    return [
        {
            "title": title,
            "content": text,
            "section_type": "Document",
        }
    ]


def _clean_chapters(chapters):
    cleaned = []

    for ch in chapters or []:
        title = (
            ch.get("title")
            or ch.get("chapter_title")
            or ch.get("heading")
            or "Chapter 1"
        )
        content = _clean_extracted_text(ch.get("content", ""))

        if not content:
            continue

        # Avoid saving tiny empty-looking chunks.
        if len(content.strip()) < 80:
            continue

        cleaned.append(
            {
                "title": str(title).strip() or "Chapter 1",
                "content": content,
                "section_type": ch.get("section_type", "Chapter"),
            }
        )

    return cleaned


def _clean_extracted_text(text: str) -> str:
    text = str(text or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _guess_single_chapter_title(text: str, filename: str) -> str:
    """
    Pick a reasonable title for a single-chapter fallback.
    Generic: works for Lecture, Chapter, Lesson, etc.
    """
    text = text or ""

    patterns = [
        r"\b(Lecture)\s+([0-9IVXLC]+)\b[^\n]*",
        r"\b(Chapter)\s+([0-9IVXLC]+)\b[^\n]*",
        r"\b(Lesson)\s+([0-9IVXLC]+)\b[^\n]*",
        r"\b(Unit)\s+([0-9IVXLC]+)\b[^\n]*",
        r"\b(Module)\s+([0-9IVXLC]+)\b[^\n]*",
        r"\b(Week)\s+([0-9IVXLC]+)\b[^\n]*",
        r"\b(Section)\s+([0-9IVXLC]+)\b[^\n]*",
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            title = match.group(0).strip()
            title = re.sub(r"\s+", " ", title)
            return title[:120]

    return extract_title_from_filename(filename) or "Chapter 1"


def extract_text_from_pdf(binary_data: bytes) -> str:
    """Extract plain text from a PDF file."""
    try:
        pdf = fitz.open(stream=binary_data, filetype="pdf")
        return "\n".join(page.get_text() for page in pdf)
    except Exception:
        logger.exception("Failed to extract text from PDF")
        return ""


def extract_text_from_docx(binary_data: bytes) -> str:
    """Extract plain text from a DOCX file."""
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(binary_data)
            tmp_path = tmp.name

        doc = docx.Document(tmp_path)
        return "\n".join(p.text for p in doc.paragraphs)

    except Exception:
        logger.exception("Failed to extract text from DOCX")
        return ""

    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


def split_into_chapters(text: str):
    """
    Split text into sections.

    Original version only matched Chapter/Lesson/Unit.
    This version also accepts Lecture/Module/Week/Section so lecture slides
    do not return chapters: [].
    """
    pattern = re.compile(
        r"\b(Chapter|Lecture|Lesson|Unit|Module|Week|Section)\s+([0-9IVXLC]+)\b",
        re.IGNORECASE,
    )
    matches = list(pattern.finditer(text or ""))

    chapters = []

    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)

        section_type = match.group(1).capitalize()
        section_number = match.group(2)
        title = f"{section_type} {section_number}"
        content = text[start:end].strip()

        if len(content) < 80:
            continue

        chapters.append(
            {
                "title": title,
                "content": content,
                "section_type": section_type,
            }
        )

    return chapters
